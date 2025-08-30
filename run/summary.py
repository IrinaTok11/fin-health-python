#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""summary.py
---------------------
Generate a cleaned "Summary" sheet in the source Excel workbook
and produce a Word report (3.2 Analysis of liquidity ratios) with
consistent corporate formatting.
"""

from __future__ import annotations

__author__  = "Irina Tokmianina"
__version__ = "0.9.0"
__license__ = "MIT"

import os
import sys
from math import isnan
from typing import Dict, List, Tuple

import pandas as pd
import openpyxl
from openpyxl import load_workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsdecls
from docx.oxml import OxmlElement, parse_xml


# ============================================================================
# Configuration
# ============================================================================

HEADER_FILL_RGB = "E5E5E5"         # Excel/Word header cell fill
DATA_FILL_RGB = "FFFFFF"           # Excel data cell fill
BORDER_RGB = "CCCCCC"              # Light gray bottom border

DEFAULT_FONT_NAME = "Calibri"
DEFAULT_FONT_SIZE_PT = 11

EXCEL_SUMMARY_SHEET_KEY = "summary"
YEARS_SHEET_KEY = "years"
PARAMS_SHEET_KEY = "parameters"
INCOME_SHEET_KEY = "income_statement"
BALANCE_SHEET_KEY = "balance_sheet"   # single expected sheet per user's workbook

WORD_OUTPUT_NAME = "3.2_analysis_of_liquidity_ratios.docx"
WORD_TITLE = "3.2 Analysis of liquidity ratios"
WORD_TABLE_CAPTION = "Table 3.2.1: Liquidity ratios, {y0}–{y2}"


# ============================================================================
# Utilities
# ============================================================================

def fail(msg: str) -> None:
    print(msg, file=sys.stderr)
    sys.exit(1)


def ensure_single_xlsx_in_cwd() -> str:
    data_dir = os.getcwd()
    files = [f for f in os.listdir(data_dir) if f.endswith(".xlsx") and not f.startswith("~$")]
    if len(files) != 1:
        fail(f"Please ensure exactly one .xlsx file is present. Found: {files}")
    return os.path.join(data_dir, files[0])


def set_calibri_font(run) -> None:
    run.font.name = DEFAULT_FONT_NAME
    rFonts = run._element.rPr.get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), DEFAULT_FONT_NAME)


def safe_float_to_str(val) -> str:
    return f"{val:.2f}" if isinstance(val, float) else str(val)


# ============================================================================
# Excel reading and calculations
# ============================================================================

def read_sheet_names(xlsx_path: str) -> List[str]:
    xls = pd.ExcelFile(xlsx_path)
    return xls.sheet_names


def resolve_sheet_name(requested_key: str, available: List[str]) -> str:
    if requested_key in available:
        return requested_key
    lowered = {s.lower(): s for s in available}
    if requested_key.lower() in lowered:
        return lowered[requested_key.lower()]
    for s in available:
        if s.lower().replace(" ", "_") == requested_key:
            return s
    raise KeyError(f"Sheet '{requested_key}' not found among: {available}")


def load_parameters(xlsx_path: str, params_sheet: str) -> Dict[str, float]:
    df = pd.read_excel(xlsx_path, sheet_name=params_sheet)
    if "Parameter" not in df.columns or "Value" not in df.columns:
        fail(f"Parameters sheet '{params_sheet}' must have 'Parameter' and 'Value' columns.")
    return dict(zip(df["Parameter"], df["Value"]))


def load_years(xlsx_path: str, years_sheet: str) -> List[int]:
    df = pd.read_excel(xlsx_path, sheet_name=years_sheet)
    if df.shape[1] < 1:
        fail(f"Years sheet '{years_sheet}' must have at least one column with years.")
    try:
        years = df.iloc[:, 0].astype(int).tolist()
    except Exception as exc:
        fail(f"Failed to parse years in sheet '{years_sheet}': {exc}")
    if len(years) < 3:
        fail("Expected at least three years (e.g., 2021–2023).")
    return years


def read_fin_table(xlsx_path: str, sheet_name: str, years: List[int]) -> pd.DataFrame:
    df = pd.read_excel(xlsx_path, sheet_name=sheet_name, header=0)
    df.columns = [str(c).strip() if not pd.isna(c) else "" for c in df.columns]
    if "Variable_Name" in df.columns and "Variable" not in df.columns:
        df = df.rename(columns={"Variable_Name": "Variable"})
    if "Variable" not in df.columns:
        fail(f"Sheet '{sheet_name}' must include 'Variable' or 'Variable_Name'.")
    year_cols = [str(y) for y in years if str(y) in df.columns]
    if not year_cols:
        fail(f"Sheet '{sheet_name}' is missing year columns among {years}. Found: {list(df.columns)}")
    df = df[["Variable"] + year_cols].copy()
    for c in year_cols:
        df[c] = pd.to_numeric(df[c], errors="coerce").fillna(0.0)
    return df.set_index("Variable")


def series_trend_and_change(series: pd.Series) -> Tuple[str, float]:
    s0, sN = series.iloc[0], series.iloc[-1]
    if s0 is None or sN is None or isnan(s0) or isnan(sN):
        return "→", 0.0
    arrow = "↑" if sN > s0 else "↓" if sN < s0 else "→"
    return arrow, round(float(sN) - float(s0), 2)


def compute_ratios(inc: pd.DataFrame, bal: pd.DataFrame, years: List[int], months_per_year: float) -> pd.DataFrame:
    """
    Compute liquidity + additional ratios used in Excel 'summary'.
    CA = inventories + trade_receivables + cash_and_equivalents
    CL = current_borrowings + trade_payables
    EBITDA margin ≈ operating_profit_loss / revenue
    """
    def get(table: pd.DataFrame, var: str, year: int) -> float:
        if var not in table.index:
            return 0.0
        return float(table.at[var, str(year)])

    recs = []
    for y in years:
        cash = get(bal, "cash_and_equivalents", y)
        inv = get(bal, "inventories", y)
        recv = get(bal, "trade_receivables", y)
        ppe = get(bal, "property_plant_equipment", y)
        curr_borr = get(bal, "current_borrowings", y)
        noncurr_borr = get(bal, "non_current_borrowings", y)
        payables = get(bal, "trade_payables", y)
        share_cap = get(bal, "share_capital", y)
        ret_earn = get(bal, "retained_earnings_loss", y)

        curr_assets = inv + recv + cash
        curr_liab = curr_borr + payables
        total_assets = ppe + curr_assets
        equity = share_cap + ret_earn
        interest_bearing_debt = curr_borr + noncurr_borr

        revenue = get(inc, "revenue", y)
        op_profit = get(inc, "operating_profit_loss", y)
        net_profit = get(inc, "profit_loss", y)

        cash_ratio = (cash / curr_liab) if curr_liab else None
        current_ratio = (curr_assets / curr_liab) if curr_liab else None
        avg_monthly_rev = (revenue / months_per_year) if months_per_year else None
        months_to_repay = (curr_liab / avg_monthly_rev) if avg_monthly_rev else None

        quick_ratio = ((cash + recv) / curr_liab) if curr_liab else None
        equity_ratio = (equity / total_assets) if total_assets else None
        debt_to_equity = (interest_bearing_debt / equity) if equity else None
        wcta = ((curr_assets - curr_liab) / curr_assets) if curr_assets else None
        roa = (net_profit / total_assets) if total_assets else None
        roe = (net_profit / equity) if equity else None
        npm = (net_profit / revenue) if revenue else None
        ebitda_m = (op_profit / revenue) if revenue else None

        recs.append({
            "year": y,
            "Cash ratio": cash_ratio,
            "Current ratio": current_ratio,
            "Quick ratio": quick_ratio,
            "Months to repay": months_to_repay,
            "Equity ratio": equity_ratio,
            "Debt to equity": debt_to_equity,
            "WCTA": wcta,
            "ROA": roa,
            "ROE": roe,
            "Net profit margin": npm,
            "EBITDA margin": ebitda_m,
        })

    df = pd.DataFrame(recs).set_index("year")

    for metric in ["Cash ratio","Current ratio","Quick ratio","Months to repay","Equity ratio","Debt to equity","WCTA","ROA","ROE","Net profit margin","EBITDA margin"]:
        arrow, diff = series_trend_and_change(df[metric])
        df[f"Trend {metric}"] = arrow
        df[f"Change {metric}"] = diff

    return df


def build_summary_table(df_ratios: pd.DataFrame, years: List[int], ratio_norms: pd.DataFrame) -> pd.DataFrame:
    rn = ratio_norms.set_index("ratio_key")
    percent_metrics = {"Equity ratio","ROA","ROE","Net profit margin","EBITDA margin"}
    metrics = ["Cash ratio", "Current ratio", "Months to repay", "Quick ratio", "Equity ratio", "Debt to equity", "WCTA", "ROA", "ROE", "Net profit margin", "EBITDA margin"]
    cols = ["Metric","Benchmark"] + [str(y) for y in years] + ["Trend", "Change"]
    out = pd.DataFrame(columns=cols)

    key_map = {"Cash ratio":"cash_ratio","Current ratio":"current_ratio","Months to repay":"months_to_repay_cl","Quick ratio":"quick_ratio","Equity ratio":"equity_ratio","Debt to equity":"debt_to_equity","WCTA":"wc_to_current_assets","ROA":"roa","ROE":"roe","Net profit margin":"net_profit_margin","EBITDA margin":"ebitda_margin"}

    for m in metrics:
        k = key_map[m]
        better = str(rn.at[k, "better_is"]).strip().lower() if k in rn.index else ""
        low = rn.at[k, "norm_low"] if k in rn.index and "norm_low" in rn.columns else None
        high = rn.at[k, "norm_high"] if k in rn.index and "norm_high" in rn.columns else None
        unit = rn.at[k, "unit"] if k in rn.index and "unit" in rn.columns else ""
        if better == "between" and pd.notna(low) and pd.notna(high):
            if str(unit).strip()=="%":
                bench = f"{low*100:g}–{high*100:g} %"
            else:
                bench = f"{low:g}–{high:g} {unit}".strip()
            bench = bench.replace(" x","")
        elif better == "higher" and pd.notna(low):
            if str(unit).strip()=="%":
                bench = f"≥{low*100:g} %"
            else:
                bench = f"≥{low:g} {unit}".strip()
            bench = bench.replace(" x","")
        elif better == "lower" and pd.notna(high):
            if str(unit).strip()=="%":
                bench = f"≤{high*100:g} %"
            else:
                bench = f"≤{high:g} {unit}".strip()
            bench = bench.replace(" x","")
        else:
            bench = ""

        row = {"Metric": m, "Benchmark": bench}
        for y in years:
            v = df_ratios.at[y, m]
            row[str(y)] = (None if v is None else float(v)) if m in percent_metrics else (f"{v:.2f}" if v is not None else "-")
        row["Trend"] = df_ratios[f"Trend {m}"].iloc[-1]
        ch = df_ratios[f"Change {m}"].iloc[-1]
        row["Change"] = float(ch) if m in percent_metrics else (f"+{ch:.2f}" if ch > 0 else f"{ch:.2f}")
        out = pd.concat([out, pd.DataFrame([row])], ignore_index=True)

    return out


# ============================================================================
# Excel writing & formatting
# ============================================================================

def write_summary_sheet(xlsx_path: str, summary_df: pd.DataFrame, summary_sheet_name: str) -> None:
    with pd.ExcelWriter(xlsx_path, engine="openpyxl", mode="a", if_sheet_exists="replace") as writer:
        summary_df.to_excel(writer, sheet_name=summary_sheet_name, index=False)

    wb = load_workbook(xlsx_path)
    ws = wb[summary_sheet_name]

    ws.column_dimensions["A"].width = 27
    for col in ["B","C","D","E","F","G","H","I","J","K","L"]:
        ws.column_dimensions[col].width = 18

    header_fill = PatternFill(start_color=HEADER_FILL_RGB, end_color=HEADER_FILL_RGB, fill_type="solid")
    for cell in ws[1]:
        cell.font = Font(name=DEFAULT_FONT_NAME, size=10, bold=True)
        cell.alignment = Alignment(horizontal="center", vertical="center")
        cell.fill = header_fill
    ws.row_dimensions[1].height = 19.2

    for row in ws.iter_rows(min_row=2, max_row=ws.max_row, min_col=1, max_col=ws.max_column):
        for idx, cell in enumerate(row):
            if idx == 0:
                cell.alignment = Alignment(horizontal="left", vertical="center")
            else:
                cell.alignment = Alignment(horizontal="center", vertical="center")
            cell.font = Font(name=DEFAULT_FONT_NAME, size=11, bold=False)
            cell.fill = PatternFill(start_color=DATA_FILL_RGB, end_color=DATA_FILL_RGB, fill_type="solid")

    thin_bottom = Border(bottom=Side(style="thin", color=BORDER_RGB))
    for row in ws.iter_rows(min_row=1, max_row=ws.max_row, min_col=1, max_col=ws.max_column):
        for cell in row:
            cell.border = thin_bottom

    # Percentage formatting: values & Change
    percent_metric_names = {"Equity ratio","ROA","ROE","Net profit margin","EBITDA margin"}
    trend_col = ws.max_column - 1
    change_col = ws.max_column
    for r in range(2, ws.max_row+1):
        metric_name = ws.cell(row=r, column=1).value
        if metric_name in percent_metric_names:
            for c in range(3, trend_col):  # only year values
                cell = ws.cell(row=r, column=c)
                if isinstance(cell.value, (int, float)):
                    cell.number_format = '0.00%'
            ch_cell = ws.cell(row=r, column=change_col)
            if isinstance(ch_cell.value, (int, float)):
                ch_cell.number_format = '+0%;-0%;0%'  # no decimals

    # Make 'summary' active on open
    try:
        wb.active = wb.sheetnames.index(summary_sheet_name)
    except Exception:
        pass

    wb.save(xlsx_path)


# ============================================================================
# Word report
# ============================================================================

def metrics_metadata(years: List[int]):
    """Descriptions, norms, and templated analysis per metric (robust on missing data)."""
    def s(x):
        try:
            return f"{float(x):.2f}"
        except Exception:
            return "-"

    return {
        "Cash ratio": {
            "title": "Cash ratio",
            "full_name": "indicates the extent to which current liabilities can be settled immediately using cash on hand.",
            "norm": "A commonly referenced interval is 0.2–0.5.",
            "analysis": lambda v: (
                f"Across {years[0]}–{years[2]}, the ratio was {s(v[0])} → {s(v[1])} → {s(v[2])}. "
                "It remained below the reference range in each year."
            ),
        },
        "Current ratio": {
            "title": "Current ratio",
            "full_name": "assesses coverage of short‑term obligations by current assets.",
            "norm": "A prudent reference range is 1–2.",
            "analysis": lambda v: (
                f"In {years[0]}–{years[1]} the ratio stood at {s(v[0])} and {s(v[1])} (around or above the range), "
                f"then fell below the range to {s(v[2])} in {years[2]}."
            ),
        },
        "Quick ratio": {
            "title": "Quick ratio",
            "full_name": "measures immediate coverage excluding inventories.",
            "norm": "Levels at or above 0.8 are generally viewed as comfortable.",
            "analysis": lambda v: (
                f"Over the period the ratio was {s(v[0])}, {s(v[1])}, and {s(v[2])}; "
                "in the final year it was below 0.8."
            ),
        },
        "Months to repay": {
            "title": "Months to repay",
            "full_name": "shows how many months of average monthly revenue would be required to repay current obligations in full.",
            "norm": "Staying under three months is generally considered prudent.",
            "analysis": lambda v: (
                f"The indicator remained within prudent limits across the period "
                f"({s(v[0])}, {s(v[1])}, {s(v[2])})."
            ),
        },
    }


def write_word_report(
    output_dir: str,
    years: List[int],
    summary_df: pd.DataFrame,
    company_name: str,
) -> str:
    df_summary = summary_df.set_index("Metric")
    out_path = os.path.join(output_dir, WORD_OUTPUT_NAME)

    meta = metrics_metadata(years)

    doc = Document()

    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)

    doc.styles["Normal"].font.name = DEFAULT_FONT_NAME
    doc.styles["Normal"].font.size = Pt(DEFAULT_FONT_SIZE_PT)

    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run_t = p_title.add_run(WORD_TITLE)
    run_t.font.name = DEFAULT_FONT_NAME
    run_t.font.size = Pt(14)
    run_t.bold = True
    run_t.font.color.rgb = RGBColor(0, 0, 0)
    set_calibri_font(run_t)
    p_title.paragraph_format.space_after = Pt(12)

    # Intro
    owner = (company_name + "’s") if company_name else "the company’s"
    p_intro = doc.add_paragraph(
        "Liquidity ratios assess a company’s ability to meet its short‑term financial obligations "
        "using readily available assets. These measures indicate whether the company can comfortably "
        "manage its short‑term liabilities with its existing resources. "
        f"Table 3.2.1 summarises {owner} key liquidity ratios for {years[0]}–{years[2]}."
    )
    p_intro.paragraph_format.space_after = Pt(15)
    p_intro.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # Caption
    p_caption = doc.add_paragraph()
    p_caption.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    caption_text = (f"Table 3.2.1: {company_name} liquidity ratios, {years[0]}–{years[2]}"
                    if company_name else WORD_TABLE_CAPTION.format(y0=years[0], y2=years[2]))
    run_c = p_caption.add_run(caption_text)
    set_calibri_font(run_c)
    run_c.font.size = Pt(DEFAULT_FONT_SIZE_PT)
    run_c.bold = True

    # Table (header + 4 rows)
    table = doc.add_table(rows=5, cols=7)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Enforce zero margins at the table level to eliminate default cell padding
    try:
        tblPr = table._tbl.tblPr
        if tblPr is None:
            tblPr = table._tbl.get_or_add_tblPr()
    except Exception:
        tblPr = table._tbl.get_or_add_tblPr()
    # Remove existing cell margins and set all to zero
    try:
        existing_tblMar = tblPr.find(qn("w:tblCellMar"))
        if existing_tblMar is not None:
            tblPr.remove(existing_tblMar)
        tblCellMar = OxmlElement("w:tblCellMar")
        for side in ("top","left","bottom","right"):
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:w"), "0")
            el.set(qn("w:type"), "dxa")
            tblCellMar.append(el)
        tblPr.append(tblCellMar)
        # Remove any table indent to avoid global left offset
        tblInd = tblPr.find(qn("w:tblInd"))
        if tblInd is not None:
            tblPr.remove(tblInd)
    except Exception:
        pass


    hdr = table.rows[0].cells
    hdr[0].text = "Metric"
    hdr[1].text = "Benchmark"
    hdr[2].text = str(years[0])
    hdr[3].text = str(years[1])
    hdr[4].text = str(years[2])
    hdr[5].text = "Trend"
    hdr[6].text = "Change"

    order = ["Cash ratio", "Current ratio", "Quick ratio", "Months to repay"]
    for i, metric in enumerate(order):
        row = table.rows[i + 1].cells
        row[0].text = metric
        row[1].text = safe_float_to_str(df_summary.loc[metric, "Benchmark"])
        row[2].text = safe_float_to_str(df_summary.loc[metric, str(years[0])])
        row[3].text = safe_float_to_str(df_summary.loc[metric, str(years[1])])
        row[4].text = safe_float_to_str(df_summary.loc[metric, str(years[2])])
        row[5].text = str(df_summary.loc[metric, "Trend"])
        row[6].text = safe_float_to_str(df_summary.loc[metric, "Change"])

    # Table styling
    for r_idx, row in enumerate(table.rows):
        # First-column exact zero padding: override cell margins and paragraph indents
        try:
            for r_i, r in enumerate(table.rows):
                for c_i, cell in enumerate(r.cells):
                    if r_i > 0 and c_i == 0:  # data rows, first column
                        # Zero paragraph indents
                        for par in cell.paragraphs:
                            try:
                                par.paragraph_format.left_indent = Pt(0)
                                par.paragraph_format.first_line_indent = Pt(0)
                            except Exception:
                                pass
                        # Zero cell-specific margins
                        try:
                            tc = cell._tc
                            tcPr = tc.get_or_add_tcPr()
                            # Remove any tcMar first
                            cur = tcPr.find(qn("w:tcMar"))
                            if cur is not None:
                                tcPr.remove(cur)
                            tcMar = OxmlElement("w:tcMar")
                            left = OxmlElement("w:left")
                            left.set(qn("w:w"), "0")
                            left.set(qn("w:type"), "dxa")
                            right = OxmlElement("w:right")
                            right.set(qn("w:w"), "0")
                            right.set(qn("w:type"), "dxa")
                            tcMar.append(left); tcMar.append(right)
                            tcPr.append(tcMar)
                        except Exception:
                            pass
        except Exception:
            pass

        for c_idx, cell in enumerate(row.cells):
            for paragraph in cell.paragraphs:
                spacing = OxmlElement("w:spacing")
                spacing.set(qn("w:line"), "240")
                spacing.set(qn("w:lineRule"), "auto")
                spacing.set(qn("w:after"), "0")
                spacing.set(qn("w:before"), "0")
                paragraph._element.get_or_add_pPr().append(spacing)
                paragraph.paragraph_format.alignment = (
                    WD_PARAGRAPH_ALIGNMENT.LEFT if (r_idx > 0 and c_idx == 0) else WD_PARAGRAPH_ALIGNMENT.CENTER
                )
                for run in paragraph.runs:
                    run.font.name = DEFAULT_FONT_NAME
                    run.font.size = Pt(DEFAULT_FONT_SIZE_PT)
                    run.bold = True if r_idx == 0 else False

    header_row = table.rows[0]
    for cell in header_row.cells:
        cell._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{HEADER_FILL_RGB}"/>'))
    for row in table.rows[1:]:
        for cell in row.cells:
            cell._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="FFFFFF"/>'))

    for r_index, row in enumerate(table.rows):
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            borders = tcPr.find(qn("w:tcBorders"))
            if borders is not None:
                tcPr.remove(borders)
            if r_index > 0:
                borders = parse_xml(
                    f'<w:tcBorders {nsdecls("w")}>'
                    f'<w:bottom w:val="single" w:sz="8" w:color="{BORDER_RGB}"/>'
                    f"</w:tcBorders>"
                )
                tcPr.append(borders)

    column_widths = [Cm(3.7), Cm(2.5), Cm(2.35), Cm(2.35), Cm(2.35), Cm(1.7), Cm(2.1)]
    for i, width in enumerate(column_widths):
        for cell in table.columns[i].cells:
            cell.width = width

    for r_idx, row in enumerate(table.rows):
        row.height = Cm(1.0 if r_idx == 0 else 0.6)
        for c_idx, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if r_idx in (1, 2, 3, 4) and c_idx == 0:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                marL = OxmlElement("w:marL")
                marL.set(qn("w:w"), "141")
                marL.set(qn("w:type"), "dxa")
                tcPr.append(marL)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # --- Prep for narrative: norms & series values used below ---
    def _as_float(v):
        try:
            s = str(v).replace('%','').replace('x','').replace('months','').strip()
            return float(s)
        except Exception:
            return float('nan')
    cash_vals  = [_as_float(df_summary.loc["Cash ratio", str(y)]) for y in years] if "Cash ratio" in df_summary.index else []
    curr_vals  = [_as_float(df_summary.loc["Current ratio", str(y)]) for y in years] if "Current ratio" in df_summary.index else []
    quick_vals = [_as_float(df_summary.loc["Quick ratio", str(y)]) for y in years] if "Quick ratio" in df_summary.index else []
    m_vals     = [_as_float(df_summary.loc["Months to repay", str(y)]) for y in years] if "Months to repay" in df_summary.index else []
    # Backward-compatible alias
    mtr_vals = m_vals

    curr_norm_low, curr_norm_high = 1.0, 2.0
    quick_norm_min = 0.8
    months_norm_max = 3.0
    mtr_norm_high = months_norm_max
    cash_norm_text = "A commonly referenced range is 0.2–0.5."
    current_norm_text = "A prudent range is 1-2."
    quick_norm_text = "A prudent minimum is 0.8."
    months_norm_text = "In practice, a sound position is to keep it under three months."
    if cash_vals:
        peak_val = max(cash_vals)
        peak_year = years[cash_vals.index(peak_val)]
    else:
        peak_val = float('nan')
        peak_year = years[0]
    # --- Narrative paragraphs (bold metric names, full third metric name) ---


    # Dynamic phrase for Cash ratio relative to 0.2–0.5 range
    cash_min, cash_max = (min(cash_vals), max(cash_vals)) if cash_vals else (float('nan'), float('nan'))
    any_below = any(x < 0.2 for x in cash_vals) if cash_vals else False
    any_above = any(x > 0.5 for x in cash_vals) if cash_vals else False
    all_below = all(x < 0.2 for x in cash_vals) if cash_vals else False
    all_within = all((0.2 <= x <= 0.5) for x in cash_vals) if cash_vals else False
    if all_below:
        cash_phrase = f"Across {years[0]}–{years[2]} the ratio remained below this range, "
    elif all_within:
        cash_phrase = f"Across {years[0]}–{years[2]} the ratio remained within this range, "
    elif (not any_below) and any_above:
        cash_phrase = f"Across {years[0]}–{years[2]} the ratio moved above the range at times, "
    elif any_below and (not any_above):
        cash_phrase = f"Across {years[0]}–{years[2]} the ratio approached the range but stayed below in some years, "
    else:
        cash_phrase = f"Across {years[0]}–{years[2]} the ratio fluctuated around the range, dipping below in some years and exceeding it in others, "

    # 1) Cash ratio paragraph
    p1 = doc.add_paragraph()
    r1 = p1.add_run("Cash ratio ")
    r1.bold = True
    p1.add_run(
        f"indicates the extent to which current liabilities can be settled immediately using cash on hand. "
        f"{cash_norm_text} " + cash_phrase +
        f"peaking at {peak_val:.2f} in {peak_year}."
    )
    p1.paragraph_format.space_after = Pt(6)
    p1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # 2) Current ratio paragraph
    v0, v1, v2 = curr_vals
    p2 = doc.add_paragraph()
    r2 = p2.add_run("Current ratio ")
    r2.bold = True
    intro_cr = "assesses the ability to cover short-term obligations with current assets. A prudent range is 1-2. "
    if v0 > curr_norm_high and v1 > curr_norm_high and v2 < curr_norm_low:
        body_cr = f"The indicator was above the benchmark in {years[0]}–{years[1]} ({v0:.2f} → {v1:.2f}) and then fell sharply below the benchmark, reaching {v2:.2f} in {years[2]}."
    elif v0 > curr_norm_high and v1 > curr_norm_high:
        body_cr = f"The indicator remained above the benchmark in {years[0]}–{years[1]} ({v0:.2f} → {v1:.2f})."
    elif v0 >= curr_norm_low and v1 >= curr_norm_low and v2 < curr_norm_low:
        body_cr = f"The indicator stayed within the prudent range in {years[0]}–{years[1]} ({v0:.2f} → {v1:.2f}) and then slipped below the benchmark to {v2:.2f} in {years[2]}."
    else:
        body_cr = f"The ratio moved from {v0:.2f} in {years[0]} to {v1:.2f} in {years[1]} and {v2:.2f} in {years[2]}."
    p2.add_run(intro_cr + body_cr)
    p2.paragraph_format.space_after = Pt(6)
    p2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # 3) Quick ratio paragraph
    vq0, vq1, vq2 = quick_vals
    p2b = doc.add_paragraph()
    r2b = p2b.add_run("Quick ratio ")
    r2b.bold = True
    intro_q = "captures the ability to meet short-term obligations from the most liquid current assets (cash and trade receivables), excluding inventories. A prudent minimum is 0.8. "
    if vq0 >= 0.8 and vq1 >= 0.8 and vq2 >= 0.8:
        body_q = f"The indicator remained at or above the prudent minimum throughout, moving from {vq0:.2f} in {years[0]} to {vq1:.2f} in {years[1]} and {vq2:.2f} in {years[2]}."
    elif vq2 < 0.8 and (vq0 >= 0.8 or vq1 >= 0.8):
        body_q = f"After touching {vq1:.2f} in {years[1]}, the indicator finished below the prudent minimum at {vq2:.2f} in {years[2]}."
    elif vq2 < 0.8:
        body_q = f"The indicator remained below the prudent minimum, changing from {vq0:.2f} to {vq1:.2f} and {vq2:.2f} by {years[2]}."
    else:
        body_q = f"The ratio moved from {vq0:.2f} in {years[0]} to {vq1:.2f} in {years[1]} and {vq2:.2f} in {years[2]}."
    p2b.add_run(intro_q + body_q)
    p2b.paragraph_format.space_after = Pt(6)
    p2b.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

# 4) Months to repay current liabilities paragraph
    m0, m1, m2 = mtr_vals
    p3 = doc.add_paragraph()
    r3 = p3.add_run("Months to repay current liabilities ")
    r3.bold = True
    intro_m = "shows how many months of average monthly revenue are required to repay current obligations in full. In practice, a sound position is to keep it under three months. "
    if all(x <= mtr_norm_high for x in [m0, m1, m2]):
        body_m = f"Over the period, the indicator remained within this benchmark, moving from {m0:.2f} months in {years[0]} to {m1:.2f} in {years[1]} and {m2:.2f} in {years[2]}."
    else:
        body_m = f"Over the period, the indicator moved from {m0:.2f} months in {years[0]} to {m1:.2f} in {years[1]} and {m2:.2f} in {years[2]}."
    p3.add_run(intro_m + body_m)
    p3.paragraph_format.space_after = Pt(12)
    p3.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

# Overall assessment + bullets
    # Build flags
    def _clean_float(x):
        s = str(x).replace("%","").replace("x","").replace("months","").strip()
        try:
            return float(s)
        except Exception:
            return float("nan")

    cash_vals  = [_clean_float(df_summary.loc["Cash ratio", str(y)]) for y in years] if "Cash ratio" in df_summary.index else []
    curr_vals  = [_clean_float(df_summary.loc["Current ratio", str(y)]) for y in years] if "Current ratio" in df_summary.index else []
    quick_vals = [_clean_float(df_summary.loc["Quick ratio", str(y)]) for y in years] if "Quick ratio" in df_summary.index else []
    m_vals     = [_clean_float(df_summary.loc["Months to repay", str(y)]) for y in years] if "Months to repay" in df_summary.index else []

    flags = []
    if cash_vals and all(v < 0.2 for v in cash_vals):
        flags.append("cash ratio persistently below the reference range")
    if curr_vals and curr_vals[-1] < 1.0:
        flags.append("current ratio dropping below 1 in the final year")
    if quick_vals and quick_vals[-1] < 0.8:
        flags.append("quick ratio below 0.8 in the final year")

    p4 = doc.add_paragraph()
    rlab = p4.add_run("Overall assessment: ")
    rlab.bold = True
    if m_vals and all(x <= 3.0 for x in m_vals):
        mitigant = "The time-to-repay metric remained within prudent limits, which partly mitigates near-term risk."
    else:
        mitigant = "No consistent mitigant from the time-to-repay metric over the period."

    if flags:
        recommendation = " Strengthening cash buffers, tightening receivables and inventory management, and aligning the maturity profile of obligations would enhance resilience."
        body_ov = "Taken together, " + (owner) + " short-term liquidity appears constrained: " + ", ".join(flags) + ". " + mitigant + recommendation
    else:
        body_ov = "Taken together, " + (owner) + " key indicators are broadly aligned with reference levels. " + mitigant

    p4.add_run(body_ov)
    p4.paragraph_format.space_after = Pt(6)
    p4.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # Key observations
    h_obs = doc.add_paragraph()
    r_obs = h_obs.add_run("Key observations:")
    r_obs.bold = True
    h_obs.paragraph_format.space_before = Pt(6)
    h_obs.paragraph_format.space_after = Pt(3)
    h_obs.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Observation bullets
    def _obs_bullet_cash():
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(0)
        p.add_run("Cash ratio persistently below the ")
        rn = p.add_run("0.2–0.5")
        rn.bold = True
        p.add_run(" reference range.")
        return p

    def _obs_bullet_current():
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(0)
        p.add_run("Current ratio dropping below ")
        rn = p.add_run("1")
        rn.bold = True
        p.add_run(" in the final year.")
        return p

    def _obs_bullet_quick():
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(0)
        p.add_run("Quick ratio below ")
        rn = p.add_run("0.8")
        rn.bold = True
        p.add_run(" in the final year.")
        return p

    has_any = False
    for f in flags:
        if "cash ratio persistently below" in f:
            _obs_bullet_cash(); has_any = True
        if "current ratio dropping below 1" in f:
            _obs_bullet_current(); has_any = True
        if "quick ratio below 0.8" in f:
            _obs_bullet_quick(); has_any = True
    if not has_any:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(0)
        p.add_run("No material weaknesses identified.")

    # Mitigant
    h_mit = doc.add_paragraph()
    r_mit = h_mit.add_run("Mitigant:")
    r_mit.bold = True
    h_mit.paragraph_format.space_before = Pt(6)
    h_mit.paragraph_format.space_after = Pt(3)
    h_mit.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    if m_vals and all(x <= 3.0 for x in m_vals):
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(0)
        p.add_run("Months to repay remained within the ")
        rn = p.add_run("3 months")
        rn.bold = True
        p.add_run(" threshold over the period.")
    else:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(0)
        p.add_run("No consistent mitigant from the time-to-repay metric (exceeded ")
        rn = p.add_run("3 months")
        rn.bold = True
        p.add_run(" in the period).")

    # Recommended actions
    h_act = doc.add_paragraph()
    r_act = h_act.add_run("Recommended actions:")
    r_act.bold = True
    h_act.paragraph_format.space_before = Pt(6)
    h_act.paragraph_format.space_after = Pt(3)
    h_act.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    for rec in [
        "Strengthen cash buffers.",
        "Tighten receivables and inventory management discipline.",
        "Align the maturity profile of obligations with projected cash flows."
    ]:
        ap = doc.add_paragraph(rec, style="List Bullet")
        ap.paragraph_format.space_after = Pt(0)

    # spacer
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    doc.save(out_path)
    return out_path


# ============================================================================
# Main
# ============================================================================

def main() -> None:
    xlsx_path = ensure_single_xlsx_in_cwd()
    sheet_names = read_sheet_names(xlsx_path)

    years_sheet = resolve_sheet_name(YEARS_SHEET_KEY, sheet_names)
    params_sheet = resolve_sheet_name(PARAMS_SHEET_KEY, sheet_names)
    income_sheet = resolve_sheet_name(INCOME_SHEET_KEY, sheet_names)
    balance_sheet = resolve_sheet_name(BALANCE_SHEET_KEY, sheet_names)
    summary_sheet = resolve_sheet_name(EXCEL_SUMMARY_SHEET_KEY, sheet_names)

    params = load_parameters(xlsx_path, params_sheet)
    months_per_year = float(params.get("months_per_year", 12))
    years = load_years(xlsx_path, years_sheet)

    inc = read_fin_table(xlsx_path, income_sheet, years)
    bal = read_fin_table(xlsx_path, balance_sheet, years)

    df_ratios = compute_ratios(inc, bal, years, months_per_year)
    ratio_norms = pd.read_excel(xlsx_path, sheet_name=resolve_sheet_name('ratio_norms', sheet_names))
    summary_df = build_summary_table(df_ratios, years, ratio_norms)

    write_summary_sheet(xlsx_path, summary_df, summary_sheet)

    output_dir = os.getcwd()
    company_name = str(params.get('company_name') or params.get('company') or params.get('Company') or params.get('company name') or '').strip()
    word_path = write_word_report(output_dir, years, summary_df, company_name)

    print("✔ Summary sheet updated and styled.")
    print(f"✔ Word report created: {word_path}")


if __name__ == "__main__":
    main()

